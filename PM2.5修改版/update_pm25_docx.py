from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_BREAK
from docx.shared import Inches
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parent
DOCX = ROOT / "PM2.5_扩充版.docx"
FIG1 = ROOT / "matlab_figures" / "fig_4_5_extension_prediction.png"
FIG2 = ROOT / "matlab_figures" / "fig_4_6_extension_control.png"


def copy_paragraph_format(src, dst):
    ppr = deepcopy(src._p.pPr) if src._p.pPr is not None else None
    if ppr is not None:
        old = dst._p.pPr
        if old is not None:
            dst._p.remove(old)
        dst._p.insert(0, ppr)


def copy_run_format(src_run, dst_run):
    rpr = deepcopy(src_run._r.rPr) if src_run is not None and src_run._r.rPr is not None else None
    if rpr is not None:
        dst_run._r.insert(0, rpr)


def set_paragraph_text(paragraph, text, template=None):
    template = template or paragraph
    src_run = template.runs[0] if template.runs else None
    ppr = deepcopy(template._p.pPr) if template._p.pPr is not None else None
    rpr = deepcopy(src_run._r.rPr) if src_run is not None and src_run._r.rPr is not None else None
    for child in list(paragraph._p):
        paragraph._p.remove(child)
    if ppr is not None:
        paragraph._p.insert(0, ppr)
    run = paragraph.add_run(text)
    if rpr is not None:
        run._r.insert(0, rpr)
    return paragraph


def insert_paragraph_after(paragraph, text, template):
    new_p = deepcopy(template._p)
    for child in list(new_p):
        if not child.tag.endswith("}pPr"):
            new_p.remove(child)
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    return set_paragraph_text(new_para, text, template)


def set_cell_text(cell, text):
    para = cell.paragraphs[0]
    template = para
    set_paragraph_text(para, text, template)
    for extra in list(cell.paragraphs[1:]):
        cell._tc.remove(extra._p)


def find_para(doc, prefix):
    for i, para in enumerate(doc.paragraphs):
        if para.style.name.lower().startswith("toc"):
            continue
        if para.text.strip().startswith(prefix):
            return i, para
    raise ValueError(f"Paragraph not found: {prefix}")


def add_picture_to_blank(paragraph, image_path, template=None, page_break_before=True):
    ppr = deepcopy(template._p.pPr) if template is not None and template._p.pPr is not None else None
    for child in list(paragraph._p):
        paragraph._p.remove(child)
    if ppr is not None:
        paragraph._p.insert(0, ppr)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    if page_break_before:
        run.add_break(WD_BREAK.PAGE)
    run.add_picture(str(image_path), width=Inches(5.6))


doc = Document(DOCX)
image_template = next(
    para for para in doc.paragraphs if para._p.xpath(".//w:drawing")
)

i46, p46 = find_para(doc, "4.6 拓展研究一")
i47, p47 = find_para(doc, "4.7 拓展研究二")
_, p5 = find_para(doc, "第5章 程序设计与分析")

body_template = doc.paragraphs[i46 + 1]
caption_template = doc.paragraphs[i46 + 3]

texts46 = [
    "在基础要求已经完成插值和最小二乘拟合的基础上，本文将第一个拓展设计为基于验证集的短期预测综合比较。预测实验仍采用留后五天检验：以前25天数据建立模型，以第26至第30天观测值作为检验集。在原有全局插值、三次样条和多项式最小二乘基础上，进一步加入指数平滑法，并把二次、三次最小二乘的正规方程改用LU分解求解。这样既能比较严格通过节点的插值模型，也能比较允许残差但追求整体趋势的拟合和平滑模型。",
    "具体地，令A为训练集上的Vandermonde矩阵，多项式最小二乘参数满足法方程A^T A a=A^T y。MATLAB程序中采用[L,U,P]=lu(A^T A)，再解U a=L^{-1}P A^T y，避免显式求逆带来的数值放大。指数平滑法采用S_t=alpha y_t+(1-alpha)S_{t-1}，取alpha=0.6，得到第25天平滑水平71.8092，并将该水平作为第26至第30天的短期预测值。",
]
set_paragraph_text(doc.paragraphs[i46 + 1], texts46[0], body_template)
set_paragraph_text(doc.paragraphs[i46 + 2], texts46[1], body_template)

table43 = doc.tables[15]
set_cell_text(table43.rows[4].cells[0], "LU三次最小二乘")
set_cell_text(table43.rows[4].cells[3], "多项式中误差最小")
new_row_xml = deepcopy(table43.rows[-1]._tr)
table43._tbl.append(new_row_xml)
new_row = table43.rows[-1]
for cell, text in zip(
    new_row.cells,
    ["指数平滑法", "7.8092", "8.3057", "验证误差最小"],
):
    set_cell_text(cell, text)

blank_fig1 = doc.paragraphs[i46 + 4]
add_picture_to_blank(blank_fig1, FIG1, image_template)
insert_paragraph_after(
    blank_fig1,
    "图4.5 MATLAB生成的短期预测曲线与Simpson累计暴露对比",
    caption_template,
)

idx, _ = find_para(doc, "由表4.3")
set_paragraph_text(
    doc.paragraphs[idx],
    "由表4.3和图4.5可知，全局插值外推的误差远大于其他方法，这是高次多项式在区间外放大的典型表现。自然三次样条在观测区间内表现较好，但用于外推时端点外误差仍然偏大。二次最小二乘拟合把后期下降趋势延伸得过快；LU三次最小二乘在多项式模型中误差较小，而指数平滑法的MAE为7.8092、RMSE为8.3057，在本验证窗口内表现最好。",
    body_template,
)
idx, _ = find_para(doc, "从数值分析角度看，该拓展")
set_paragraph_text(
    doc.paragraphs[idx],
    "中心差分法进一步揭示了各模型对尾段变化率的刻画差异。第26至第30天观测序列的中心差分平均斜率约为-2.0000，而二次和三次最小二乘外推分别约为-5.2831和-4.6047，说明多项式模型低估了后期浓度水平；指数平滑法的斜率为0，虽然不能反映继续下降的速度，但在短期预警中给出了较保守的浓度估计。Simpson积分法计算的验证期累计暴露中，观测值为256.0000，二次和三次最小二乘分别为211.7332和221.1745，指数平滑为287.2370。由此可见，短期预测不能只看曲线是否光滑，还应结合验证误差、变化率和累计暴露共同判断。",
    body_template,
)

i47, p47 = find_para(doc, "4.7 拓展研究二")
caption_template_47 = doc.paragraphs[i47 + 3]
body_template_47 = doc.paragraphs[i47 + 1]

set_paragraph_text(
    doc.paragraphs[i47 + 1],
    "第二个自主拓展点选择达标约束下的成本最小管控方案。在基础要求中，非线性方程只求出使月均浓度降至35微克每立方米的最小管控强度。为了进一步体现数值求根和优化思想，本文令f(k)=82.5(1-0.12k)-35，并在二分法、牛顿法之外加入割线法。割线法迭代格式为k_{m+1}=k_m-f(k_m)(k_m-k_{m-1})/[f(k_m)-f(k_{m-1})]，只需要函数值，不需要显式导数，适合用于导数不易获得的管控模型。",
    body_template_47,
)
set_paragraph_text(
    doc.paragraphs[i47 + 2],
    "MATLAB中取初值k_0=3、k_1=5.5进行割线迭代，得到k=4.7979797980，与二分法和牛顿法结果一致。此时缩放因子为0.4242，管控后月均浓度为35.0000微克每立方米；若成本函数为C(k)=50k^2，则对应治理成本为1151.03万元。为了补充平均浓度指标，本文还用Simpson积分法对样条曲线下的月累计污染暴露进行估计，管控前为2407.9144，管控后为1021.5394，说明该强度不仅满足均值约束，也显著降低了整月暴露面积。",
    body_template_47,
)

blank_fig2 = doc.paragraphs[i47 + 4]
add_picture_to_blank(blank_fig2, FIG2, image_template)
insert_paragraph_after(
    blank_fig2,
    "图4.6 MATLAB生成的达标成本关系与割线法收敛过程",
    caption_template_47,
)

idx, _ = find_para(doc, "表4.4显示")
set_paragraph_text(
    doc.paragraphs[idx],
    "表4.4和图4.6显示，目标浓度每降低5微克每立方米，所需管控强度和治理成本都会明显上升。当目标由40降至35微克每立方米时，成本从921.46万元增加到1151.03万元，增加约229.57万元；若继续要求降至33微克每立方米，成本达到1250.00万元。中心差分法给出的局部灵敏度也支持这一判断：在k=4.7980附近，月均浓度对k的变化率约为-9.9000，而成本对k的变化率约为479.7980万元/单位强度，说明接近达标边界后继续加大管控会迅速增加经济负担。",
    body_template_47,
)
idx, _ = find_para(doc, "因此，本文给出的最优管控方案")
set_paragraph_text(
    doc.paragraphs[idx],
    "因此，本文给出的最优管控方案不是强度越大越好，而是在达标约束下选择最低成本强度。该拓展把割线法求根、Simpson积分评价、中心差分灵敏度和成本函数联系起来：求根确定可行边界，积分衡量累计暴露，差分刻画边际变化，成本函数决定边界上的最优点。若实际治理中还需考虑健康收益、企业限产损失和多污染物协同控制，则可以把单目标成本最小问题进一步扩展为多目标优化模型。",
    body_template_47,
)

doc.save(DOCX)
print(DOCX)
